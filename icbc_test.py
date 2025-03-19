import time
import os
import requests
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document


IMAGE_NOT_DOWNLOAD= "https://images.ctfassets.net/1eftmbczj7w9/4lxJsZKPeDdqvWh9imZdNL/ac51d64ed9eab0cb6075423af8f2cc44/VPLOGO.jpg"
images_dir = "images_download"
os.makedirs(images_dir, exist_ok=True)

# 配置 Selenium 选项
options = Options()
options.add_argument("--headless")  # 无头模式
options.add_argument("--disable-gpu")
options.add_argument("--no-sandbox")

# 启动 WebDriver
service = Service(ChromeDriverManager().install())
driver = webdriver.Chrome(service=service, options=options)
wait = WebDriverWait(driver, 10)  # 设置等待时间

# 访问 ICBC 练习考试页面
url = "https://practicetest.icbc.com/"
driver.get(url)
time.sleep(3)  # 等待加载

# **Step 1: 选择“简体中文”并确认**
try:
    simplified_chinese = wait.until(
        EC.element_to_be_clickable((By.XPATH, "//label[contains(text(), '简体中文')]"))
    )
    simplified_chinese.click()
    time.sleep(0.5)

    confirm_button = driver.find_element(By.XPATH, "//button[contains(text(), '确认')]")
    confirm_button.click()
    print("✅ Step1:语言已选择：简体中文")
    time.sleep(2)  # 进入下一个页面

except Exception as e:
    print("❌ Step1:语言选择失败:", e)
    driver.quit()

# **Step 2: 点击 "笔试练习"**
try:
    practice_test_button = wait.until(
        EC.element_to_be_clickable((By.XPATH, "//a[contains(text(), '笔试练习')]"))
    )
    practice_test_button.click()
    print("✅ Step2:进入页面2：笔试练习")
    time.sleep(2)

except Exception as e:
    print("❌ Step2:进入笔试练习失败:", e)
    driver.quit()

# **Step 3: 点击 "完整测试"**
try:
    full_test_button = wait.until(
        EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), '完整测试')]"))
    )
    full_test_button.click()
    print("✅ Step3:进入页面3：完整测试")
    time.sleep(2)

except Exception as e:
    print("❌ Step3:进入完整测试失败:", e)
    driver.quit()

# **Step 4: 开始抓取测试题**
document = Document()  # 创建 Word 文档
question_data = []  # 存储所有题目信息

while True:
    try:
        # **抓取题目文本**
        question_text = wait.until(
            EC.presence_of_element_located((By.XPATH, "//p[contains(@class, 'mb-2') and contains(@class, 'font-headings') and contains(@class, 'text-[18px]') and contains(@class, 'font-bold')]"))
        ).text
        print(f"📌 题目: {question_text}")

        # **抓取图片**
        image_url = ""
        try:
            # 定位图片
            image_element = driver.find_element(By.CSS_SELECTOR, "img.max-w-60")  # 使用 class 精准定位
            image_url = image_element.get_attribute("src")

            if IMAGE_NOT_DOWNLOAD == image_url:
                print("⚠️ 图片 不下载")
            else:
                # 确保 URL 有效
                if image_url.startswith("http") :
                    image_filename = os.path.basename(image_url).split("?")[0]  # 去掉 URL 参数，获取文件名
                    image_download_filename = os.path.join(images_dir, image_filename)
                    img_data = requests.get(image_url).content

                    # 保存图片
                    with open(image_download_filename, "wb") as f:
                        f.write(img_data)
                    print(f"✅ 图片已下载: {image_download_filename}")
                else:
                    print("⚠️ 图片 URL 无效")
        except Exception as e:
            print(f"📌 该题目无图片, 或者图片加载失败: {e}")

        # # **抓取选项**
        # options_data = []
        # option_buttons = driver.find_elements(By.TAG_NAME, "button")
        # for button in option_buttons:
        #     option_letter = button.find_element(By.CLASS_NAME, "font-bold").text  # 选项字母
        #     option_text = button.find_elements(By.TAG_NAME, "div")[1].text  # 选项内容
        #     options_data.append({"letter": option_letter, "text": option_text})

        # print(f"📌 选项: {options_data}")

        # **抓取选项**
        options_data = []
        option_buttons = driver.find_elements(By.TAG_NAME, "button")  # 找到所有按钮

        for button in option_buttons:
            divs = button.find_elements(By.TAG_NAME, "div")  # 找到按钮下所有 div
            
            if len(divs) >= 3:  # 确保有足够的 div
                option_letter = divs[1].text.strip()  # 选项字母（位于第二个 div）
                option_text = divs[2].text.strip()  # 选项内容（位于第三个 div）
                
                options_data.append({"letter": option_letter, "text": option_text})

        print(f"📌 选项: {options_data}")

        # **点击 A 选项**
        for btn in option_buttons:
            if btn.text.startswith("A"):
                btn.click()
                print("✅ 选择 A")
                break

        # **点击 "提交答案"**
        submit_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.CSS_SELECTOR, "button[type='submit'].bg-purple"))
        )
        submit_button.click()
        print("✅ 点击 提交答案")
        time.sleep(2)

        # **获取正确答案**
        try:
            # **查找 "正确" 标志，判断 A 是否正确**
            driver.find_element(By.XPATH, "//button[@value='A']//img[contains(@src, 'icon-checkmark.svg')]")
            correct_answer = "A"
        except:
            # **如果 A 错误，查找正确答案的选项字母**
            correct_answer = driver.find_element(By.XPATH, "//button[contains(@class, 'border-[#3adda2]')]//div[contains(@class, 'h-7') and contains(@class, 'w-7')]").text.strip()
        print(f"✅ 正确答案: {correct_answer}")

        # **存储题目、图片、答案**
        question_data.append({
            "question": question_text,
            "options": options_data,
            "image": image_url,
            "correct_answer": correct_answer
        })

        # **点击 "下一个问题"**
        try:
            next_button = wait.until(
                EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), '下一个问题')]"))
            )
            next_button.click()
            time.sleep(2)
            print(f"✅ 下一个问题: {correct_answer}")
        except:
            next_button = wait.until(
                EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), '完成')]"))
            )
            next_button.click()
            time.sleep(2)
            print(f"✅ 完成: ")
            break

        # break  # 退出循环
    except Exception as e:
        print("❌ 题目元素 抓取失败:", e)
        break  # 退出循环

# **Step 5: 保存到 Word 文档**
document.add_heading("ICBC 题库", level=1)

for item in question_data:
    document.add_heading(item["question"], level=2)
    for option in item["options"]:
        document.add_paragraph(f"{option['letter']}: {option['text']}")
    document.add_paragraph(f"✅ 正确答案: {item['correct_answer']}")
    
    # 添加图片
    if item["image"]:
        img_filename = os.path.basename(item["image"]).split("?")[0]
        if os.path.exists(img_filename):
            document.add_picture(img_filename)

document.save("ICBC_题库.docx")
print("✅ 所有题目已保存到 ICBC_题库.docx")

# 关闭浏览器
driver.quit()
