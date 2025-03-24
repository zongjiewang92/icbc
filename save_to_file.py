import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TM = time.strftime("%Y%m%d-%H%M%S")
# 创建结果目录
FINAL_PATH = os.path.join("final_question", TM)
os.makedirs(FINAL_PATH, exist_ok=True)


def add_page_break(document):
    """添加分页符"""
    document.add_page_break()

def save_to_word(question_data, filename_input="ICBC_题库.docx"):

    filename = os.path.join(FINAL_PATH, filename_input)
    
    """先存无图片题目，再存有图片题目，每页不固定题目数量"""
    document = Document()
    document.add_heading("ICBC 题库", level=1)
    # 获取当前时间
    current_time = datetime.now().strftime("%Y-%m-%d")  # 格式化时间
    # 添加当前时间
    document.add_heading(f"更新时间: {current_time} , 题库有效时间1个月", level=2)

    # **分离无图片和有图片的题目**
    no_image_questions = [q for q in question_data if not q["image"]]
    image_questions = [q for q in question_data if q["image"]]

    def add_questions_to_document(questions, start_index=0):
        """按顺序存储题目，每页自动分页"""
        for item in questions:
            # **题目**
            p = document.add_paragraph()
            run = p.add_run(f"Q{questions.index(item) + start_index + 1}: {item['question']}")
            run.bold = True
            run.font.size = Pt(9)  # 字体变小
            p.paragraph_format.space_after = Pt(2)  # 行间距缩小

            # **图片**
            if item["image"] and os.path.exists(item["image"]):
                p = document.add_paragraph()
                p.add_run().add_picture(item["image"], width=Inches(2.2))  # 图片缩小

            # **选项**
            for option in item["options"]:
                p = document.add_paragraph(f"{option['letter']}: {option['text']}")
                p.paragraph_format.space_after = Pt(2)  # 行间距缩小
                p.runs[0].font.size = Pt(9)  # 字体变小

            # **正确答案**
            p = document.add_paragraph(f"✅ 正确答案: {item['correct_answer']}")
            p.runs[0].font.size = Pt(9)  # 字体变小
            p.paragraph_format.space_after = Pt(2)  # 行间距缩小

            
    # **先存无图片的题目**
    add_questions_to_document(no_image_questions, 0)
    
    # **再存有图片的题目**
    add_questions_to_document(image_questions, start_index=len(no_image_questions) )

    document.save(filename)
    print(f"✅ 所有题目已保存到 {filename}")
